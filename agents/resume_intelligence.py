"""
Resume Intelligence Layer — Phase 1 (Refined).

Changes in this revision:
  - Composite organizational scale scoring (2000+ employees → enterprise)
  - Industry detection v2: contextual matching, false-positive suppression
  - Executive Framing Analysis: operational vs strategic language ratio
  - Strategic Visibility improved: board, P&L, transformation ownership signals
  - Weak visibility areas: more actionable, executive-oriented
  - LLM prompt extended: executive_framing_level + reframing_opportunities
  - Explainability: scale_rationale, industry_evidence, framing signals
"""
from __future__ import annotations

import json
import os
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Enums
# ---------------------------------------------------------------------------

class PositioningLevel(str, Enum):
    LOW    = "LOW"
    MEDIUM = "MEDIUM"
    HIGH   = "HIGH"


# ---------------------------------------------------------------------------
# Resume Profile Schema
# ---------------------------------------------------------------------------

@dataclass
class ResumeProfile:
    """
    Structured intelligence profile extracted from an executive CV.
    Single source of truth for all downstream resume-related layers.
    """
    # Source
    raw_text:    str
    source_file: str = ""
    source_type: str = ""   # txt / md / docx / pdf

    # Extracted structured data
    full_name:              str = ""
    roles_history:          list[dict[str, str]] = field(default_factory=list)
    years_experience:       int = 0
    industries:             list[str] = field(default_factory=list)

    # Intelligence signals (rule-based)
    ai_signals:             list[str] = field(default_factory=list)
    governance_signals:     list[str] = field(default_factory=list)
    transformation_experience: list[str] = field(default_factory=list)
    technology_domains:     list[str] = field(default_factory=list)
    operational_scope:      list[str] = field(default_factory=list)
    innovation_signals:     list[str] = field(default_factory=list)
    business_signals:       list[str] = field(default_factory=list)

    # Leadership & scale (composite scoring)
    leadership_scope:              str = ""   # individual / team / division / company / group
    organizational_scale:          str = ""   # startup / mid / enterprise / holding
    organizational_scale_signals:  list[str] = field(default_factory=list)
    organizational_scale_rationale: str = ""

    # Executive Framing Analysis (rule-based)
    executive_framing_level:       str = ""   # operational / technical / managerial / executive / strategic
    executive_framing_score:       int = 0    # 0–100
    operational_language_signals:  list[str] = field(default_factory=list)
    executive_language_signals:    list[str] = field(default_factory=list)

    # LLM-enriched intelligence
    executive_summary:             str = ""
    strategic_positioning:         str = ""
    professional_archetype:        str = ""
    trajectory_markers:            list[str] = field(default_factory=list)
    strongest_assets:              list[str] = field(default_factory=list)
    weak_visibility_areas:         list[str] = field(default_factory=list)
    executive_reframing_opportunities: list[dict] = field(default_factory=list)

    # Dashboard positioning levels
    ai_positioning_level:             PositioningLevel = PositioningLevel.LOW
    transformation_positioning_level: PositioningLevel = PositioningLevel.LOW
    executive_positioning_level:      PositioningLevel = PositioningLevel.LOW
    operational_depth_level:          PositioningLevel = PositioningLevel.LOW
    strategic_visibility_level:       PositioningLevel = PositioningLevel.LOW

    # Evidence references
    evidence: dict[str, list[str]] = field(default_factory=dict)

    # Meta
    analysis_timestamp: datetime = field(default_factory=datetime.now)
    confidence_notes:   list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def analyze_resume(text: str, source_file: str = "", source_type: str = "text") -> ResumeProfile:
    """Entry point. Rule-based extraction followed by LLM enrichment."""
    cleaned = _normalize_text(text)
    profile = _extract_rule_based(cleaned, source_file, source_type)
    _enrich_with_llm(profile, cleaned)
    return profile


def load_resume(path: Path) -> tuple[str, str]:
    """Load resume text. Returns (text, source_type). Supports .txt .md .docx .pdf"""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace"), suffix.lstrip(".")
    if suffix == ".docx":
        return _from_docx(path), "docx"
    if suffix == ".pdf":
        return _from_pdf(path), "pdf"
    raise ValueError(f"Unsupported file type: {suffix}. Use .txt .md .docx .pdf")


# ---------------------------------------------------------------------------
# File adapters
# ---------------------------------------------------------------------------

def _from_docx(path: Path) -> str:
    """
    Extract DOCX as Markdown, preserving document structure.
    Headings → #/##/###, bold short lines → **bold**, lists → - item,
    tables → Markdown table rows. Elements are output in document order.
    """
    try:
        import docx as _docx
        from docx.oxml.ns import qn

        doc = _docx.Document(str(path))
        lines: list[str] = []

        # Heading style names (EN + RU variants)
        _H = {
            "Heading 1": "#", "Heading 2": "##", "Heading 3": "###",
            "Заголовок 1": "#", "Заголовок 2": "##", "Заголовок 3": "###",
            "Title": "#", "Subtitle": "##",
        }
        _LIST = {"List Paragraph", "List Bullet", "List Number",
                 "Список абзацев", "Маркированный список", "Нумерованный список"}

        def _para_to_md(para) -> str | None:
            text = para.text.strip()
            if not text:
                return None
            sname = para.style.name if para.style else ""
            if sname in _H:
                return f"{_H[sname]} {text}"
            if sname in _LIST or sname.lower().startswith("list"):
                return f"- {text}"
            # Detect bold-only short lines as implicit headings
            runs = [r for r in para.runs if r.text.strip()]
            if runs and len(text) <= 80 and all(r.bold for r in runs):
                return f"**{text}**"
            return text

        def _table_to_md(table) -> list[str]:
            rows = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                # Deduplicate merged cells (python-docx repeats merged cell text)
                deduped: list[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                if any(deduped):
                    rows.append("| " + " | ".join(deduped) + " |")
            return rows

        # Iterate body elements in document order (paragraphs and tables interleaved)
        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                # Find matching paragraph object
                for para in doc.paragraphs:
                    if para._element is child:
                        md = _para_to_md(para)
                        if md:
                            lines.append(md)
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is child:
                        lines.extend(_table_to_md(table))
                        break

        return "\n".join(lines)
    except Exception as exc:
        raise ValueError(f"DOCX read error: {exc}") from exc


def _from_pdf(path: Path) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        return "\n\n".join(p.extract_text() for p in reader.pages if p.extract_text())
    except Exception as exc:
        raise ValueError(f"PDF read error: {exc}") from exc


# ---------------------------------------------------------------------------
# Text normalization
# ---------------------------------------------------------------------------

def _normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\xa0", " ").replace("­", "").replace("​", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(ln.strip() for ln in text.splitlines()).strip()


# ---------------------------------------------------------------------------
# Roles history parser
# ---------------------------------------------------------------------------

_YEAR_RANGE = re.compile(
    r"(19|20)\d{2}\s*[–—\-]\s*((19|20)\d{2}|н\.в\.|н/в|наст\.?|по\s+наст|по\s+н|present|current|сейчас)"
    r"|с\s+(19|20)\d{2}\s*(?:г\.?|года?)?\s*(?:по\s+наст|по\s+н|н\.в\.|н/в|—|–|-|$)"
    r"|(19|20)\d{2}\s*(?:г\.?\s*)?[–—\-]\s*(?:наст|н\.в\.|н/в|по\s+наст|сейчас|present)",
    re.IGNORECASE,
)
_YEAR_SINGLE = re.compile(r"\b(19|20)\d{2}\b")


def parse_roles_history(text: str) -> list[dict]:
    """
    Extract structured work positions from CV text.
    Returns list of {period, company, role, highlights}.
    Works heuristically: looks for year-range markers as role boundaries.
    """
    lines = text.splitlines()
    role_starts: list[int] = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        # Short line (header) with a year range → new role starts here
        if _YEAR_RANGE.search(stripped) and len(stripped) < 200:
            role_starts.append(i)

    if not role_starts:
        return []

    roles: list[dict] = []
    for idx, start in enumerate(role_starts):
        end = role_starts[idx + 1] if idx + 1 < len(role_starts) else len(lines)
        block = [ln.strip() for ln in lines[start:end] if ln.strip()]
        if not block:
            continue

        # Extract period from first line
        period_match = _YEAR_RANGE.search(block[0])
        period = period_match.group(0).strip() if period_match else ""

        # First line may be "period | Role" or just period; rest are company/role
        first_line = block[0]
        parts = re.split(r"\s*[|/]\s*", first_line)
        role_in_header = parts[1].strip() if len(parts) > 1 else ""
        company_in_header = parts[2].strip() if len(parts) > 2 else ""

        # Next non-empty lines after the period line
        rest = block[1:6]
        company = company_in_header or (rest[0] if rest else "")
        role    = role_in_header    or (rest[1] if len(rest) > 1 else "")

        # Bullet highlights — only explicit bullet lines (not company/header lines)
        highlights = [
            ln.lstrip("—-–•► ").strip()
            for ln in block[1:]
            if re.match(r"^[—\-–•►]", ln) and len(ln) > 10
        ][:5]

        roles.append({
            "period":     period,
            "company":    company,
            "role":       role,
            "highlights": highlights,
        })

    return roles


def format_roles_for_prompt(roles: list[dict], max_roles: int = 6) -> str:
    """Format roles history into a compact prompt-friendly string."""
    if not roles:
        return "—"
    lines: list[str] = []
    for i, r in enumerate(roles[:max_roles], 1):
        company_period = f"{r['company']} ({r['period']})" if r['period'] else r['company']
        lines.append(f"{i}. {r['role']} | {company_period}")
        for h in r['highlights'][:3]:
            if h:
                lines.append(f"   • {h[:120]}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Signal patterns
# ---------------------------------------------------------------------------

_AI_PATTERNS = [
    r"\bAI\b", r"\bML\b", r"\bLLM\b", r"\bGPT\b",
    r"machine learning", r"deep learning", r"нейросет",
    r"computer vision", r"\bNLP\b", r"MLOps", r"DataOps",
    r"AI governance", r"AI strategy", r"AI assistant",
    r"multi.?agent", r"predictive analytics", r"генеративн",
    r"decision intelligence", r"observabilit", r"автоматизац.*AI",
    r"ИИ\b", r"искусственн.*интеллект",
]

_TRANSFORM_PATTERNS = [
    r"цифров.*трансформац", r"digital transformation",
    r"цифровизац", r"digitali[sz]ation",
    r"change management", r"agile transformation",
    r"ERP.*внедр", r"внедр.*ERP",
    r"операционн.*трансформ", r"стратег.*трансформ",
    r"бизнес.трансформ",
]

_GOVERNANCE_PATTERNS = [
    r"ITSM", r"ITIL", r"COBIT", r"ISO\s*\d{4,5}",
    r"кибербезопасност", r"information security",
    r"compliance", r"риск.*управлен", r"risk management",
    r"audit", r"SLA", r"governance",
    r"ИБ\b", r"безопасност.*данных",
]

_TECH_PATTERNS = {
    "Cloud / DevOps":   [r"AWS", r"Azure", r"GCP", r"cloud", r"DevOps", r"Kubernetes", r"Docker"],
    "ERP / 1C":         [r"1С", r"1C\b", r"SAP\b", r"ERP", r"Oracle EBS"],
    "Data / Analytics": [r"BI\b", r"DWH", r"data warehouse", r"аналитик", r"Power BI", r"Tableau", r"clickhouse"],
    "Infrastructure":   [r"инфраструктур", r"сеть", r"ITSM", r"datacenter", r"ЦОД"],
    "AI / ML":          [r"\bAI\b", r"\bML\b", r"machine learning", r"нейросет", r"LLM"],
    "Security":         [r"кибербезопасност", r"ИБ\b", r"SOC\b", r"SIEM", r"pentest"],
    "Architecture":     [r"архитектур", r"микросервис", r"API", r"интеграц"],
    "Process / BPMN":   [r"BPMN", r"процессн", r"автоматизац.*процесс"],
}

# Industry patterns — contextual matching applied separately
_INDUSTRY_PATTERNS = {
    "Производство":  [r"производств", r"завод", r"фабрик", r"промышленност", r"manufacturing"],
    "Банки/Финансы": [r"банк[а-яё\s]", r"финанс.*компани", r"финанс.*холдинг", r"страхов.*компани", r"fintech\b", r"banking"],
    "Ритейл":        [r"ритейл", r"retail\b", r"торговая\s+сеть", r"e-commerce"],
    "Логистика":     [r"логистик", r"транспортная\s+компани", r"склад.*компани", r"supply chain"],
    "ИТ-компания":   [r"ИТ.компани", r"software\s+company", r"разработк.*ПО.*компани", r"tech company"],
    "Телеком":       [r"телеком.*компани", r"telecom\s+company", r"оператор\s+связи", r"mobile\s+operator"],
    "Госсектор":     [r"государствен.*компани", r"федеральн.*предприяти", r"министерств", r"госкорпорац"],
    "Консалтинг":    [r"консалтинговая\s+компани", r"consulting\s+firm", r"advisory\s+firm"],
    "Энергетика":    [r"нефтяная\s+компани", r"нефтегаз", r"энергетическая\s+компани", r"oil\s+and\s+gas"],
}

# Minimum confirmed matches to declare an industry (higher = less false positives)
_INDUSTRY_MIN_MATCHES = {
    "Производство": 1, "Логистика": 1, "Госсектор": 1, "Энергетика": 1,
    "Банки/Финансы": 1, "Ритейл": 1, "ИТ-компания": 1,
    "Телеком": 1, "Консалтинг": 1,
}

# Context words that indicate the user is a VENDOR/SUPPLIER, not an employee
_INDUSTRY_VENDOR_CONTEXT = [
    r"для\s+клиент", r"заказчик", r"для\s+банк", r"поставщик\s+для",
    r"проект.*для", r"внедрен.*для", r"обслуживан.*клиент",
]

_OPERATIONAL_SCOPE_PATTERNS = [
    r"команд.*\d+", r"\d+.*человек", r"бюджет.*\d", r"\d.*млн", r"\d.*млрд",
    r"портфел.*проект", r"P&L", r"CAPEX", r"OPEX",
]

# Strategic visibility — extended set incl. board/transformation/portfolio
_STRATEGIC_VISIBILITY_PATTERNS = [
    r"ROI", r"снижен.*затрат", r"экономи.*\d", r"\d+%", r"рост.*выруч",
    r"оптимизац.*процесс", r"KPI", r"measurable",
    r"совет\s*директор", r"board\b", r"акционер",
    r"P&L", r"бюджет.*\d+\s*млн", r"CAPEX", r"OPEX",
    r"портфел.*проект", r"кросс.функционал",
    r"операционная\s*модель", r"владел.*стратег",
]

_INNOVATION_PATTERNS = [
    r"инновац", r"R&D", r"стартап", r"MVP", r"пилот.*проект",
    r"proof of concept", r"PoC", r"новое направлен",
]

# Executive framing language patterns
_OPERATIONAL_LANG_PATTERNS = [
    r"сопровождение\b", r"поддержка\s+пользователей", r"техническая\s+поддержка",
    r"администрирование\b", r"настройка\s+оборудования", r"установка\s+ПО",
    r"help\s*desk", r"обслуживание\s+ИТ", r"мониторинг\s+серверов",
    r"uptime\b", r"устранение\s+неисправностей", r"замена\s+оборудования",
]

_EXECUTIVE_LANG_PATTERNS = [
    r"стратегия\b", r"трансформация\b", r"roadmap\b", r"governance\b",
    r"совет\s*директор", r"\bCEO\b", r"P&L\b",
    r"бюджетирование\b", r"портфель\s*проект", r"\bROI\b", r"business\s*impact",
    r"операционная\s*модель", r"организационное\s*развитие",
    r"стратегическое\s*планирование", r"cross.functional",
    r"непрерывность\s*бизнес", r"ownership.*трансформ", r"digital\s*agenda",
]


# ---------------------------------------------------------------------------
# Core helpers
# ---------------------------------------------------------------------------

def _find_signals(patterns: list[str], text: str, max_evidence: int = 3) -> tuple[list[str], list[str]]:
    found_labels: list[str] = []
    evidence: list[str] = []
    lower = text.lower()
    for pat in patterns:
        for m in re.finditer(pat, lower, re.IGNORECASE):
            label = m.group(0).strip()
            if label and label not in found_labels:
                found_labels.append(label)
            if len(evidence) < max_evidence:
                start = max(0, m.start() - 40)
                end   = min(len(text), m.end() + 60)
                fragment = "…" + text[start:end].replace("\n", " ").strip() + "…"
                if fragment not in evidence:
                    evidence.append(fragment)
    return found_labels, evidence


def _estimate_years(text: str) -> int:
    matches = re.findall(r"(19|20)(\d{2})", text)
    if matches:
        all_years = [int(f"{a}{b}") for a, b in matches]
        current = datetime.now().year
        all_years = [y for y in all_years if 1990 <= y <= current]
        if len(all_years) >= 2:
            return max(all_years) - min(all_years)
    return 0


# ---------------------------------------------------------------------------
# Composite scale detection
# ---------------------------------------------------------------------------

def _detect_scale_composite(text: str) -> tuple[str, list[str], str]:
    """
    Returns (scale, signals, rationale).
    Scale: startup / mid / enterprise / holding

    Scoring:
      holding ≥ 6 pts · enterprise ≥ 3 pts · mid ≥ 1 pt
    """
    lower = text.lower()
    score = 0
    signals: list[str] = []

    # Employee count (take the highest tier that matches)
    emp_tiers = [
        (r"(1[0-9]|[2-9]\d)[\s,.]?000\s*(сотрудник|человек|работник|staff)", 4, "10,000+ employees"),
        (r"[5-9][\s,.]?000\s*(сотрудник|человек|работник)",                  3, "5,000+ employees"),
        (r"[2-4][\s,.]?000\s*(сотрудник|человек|работник)",                  2, "2,000+ employees"),
        (r"\b[5-9]\d{2}\s*(сотрудник|человек|работник)",                     1, "500+ employees"),
        (r"\d{4,}\+?\s*(employees|сотрудников|человек)",                     2, "1,000+ employees"),
    ]
    for pat, pts, label in emp_tiers:
        if re.search(pat, lower, re.IGNORECASE):
            score += pts
            signals.append(label)
            break

    # Holding / group structure
    if re.search(r"холдинг|группа\s+компаний|group\s+of\s+companies", lower, re.IGNORECASE):
        score += 2
        signals.append("holding / group structure")

    # Multi-site / distributed
    if re.search(r"(несколько|распределён|multi.?site|территориальн|площадок\b|подразделений)", lower, re.IGNORECASE):
        score += 1
        signals.append("multi-site / distributed")

    # PMO / portfolio management
    if re.search(r"(PMO\b|проектный\s+офис|портфель\s+проект)", lower, re.IGNORECASE):
        score += 1
        signals.append("PMO / project portfolio")

    # Complex ERP landscape
    erp_set = set(re.findall(r"\b(1С|1C|SAP|Oracle\s*EBS|ERP\b|МЭС)\b", text, re.IGNORECASE))
    if len(erp_set) >= 3:
        score += 1
        signals.append(f"complex ERP landscape ({len(erp_set)} systems)")

    # Board / governance level
    if re.search(r"(COBIT\b|совет\s+директор|аудит\s+ИТ|compliance\b|board\b)", lower, re.IGNORECASE):
        score += 1
        signals.append("governance / board level")

    # Critical / production infrastructure
    if re.search(r"(критичн.*инфраструктур|production\s+uptime|ЦОД\b|datacenter|disaster\s+recovery)", lower, re.IGNORECASE):
        score += 1
        signals.append("critical infrastructure")

    # Federal / national scale
    if re.search(r"(федеральн.*масштаб|национальн.*масштаб|federal\s+scale)", lower, re.IGNORECASE):
        score += 1
        signals.append("federal / national scale")

    if score >= 6:
        scale = "holding"
    elif score >= 3:
        scale = "enterprise"
    elif score >= 1:
        scale = "mid"
    else:
        scale = "startup"

    rationale = f"{scale.upper()} (score {score}/10)"
    if signals:
        rationale += ": " + " · ".join(signals)

    return scale, signals, rationale


def _detect_leadership_scope(text: str) -> str:
    lower = text.lower()
    if re.search(r"директор|вице-президент|CIO|CTO|CDO|CDTO|CAIO|CPO|руководитель\s+\w+", lower, re.IGNORECASE):
        return "company"
    if re.search(r"руководитель\s+направлен|department head|division\s+head", lower, re.IGNORECASE):
        return "division"
    if re.search(r"руководитель\s+отдел|начальник\s+отдел|head of\b|тимлид\b|lead\b", lower, re.IGNORECASE):
        return "team"
    return "individual"


# ---------------------------------------------------------------------------
# Industry detection v2 — contextual, false-positive suppressed
# ---------------------------------------------------------------------------

def _detect_industries_v2(text: str) -> tuple[list[str], dict[str, list[str]]]:
    lower = text.lower()
    industries: list[str] = []
    industry_evidence: dict[str, list[str]] = {}

    for industry, patterns in _INDUSTRY_PATTERNS.items():
        min_matches = _INDUSTRY_MIN_MATCHES.get(industry, 1)
        confirmed: list[str] = []
        fragments: list[str] = []

        for pat in patterns:
            for m in re.finditer(pat, lower, re.IGNORECASE):
                ctx_start = max(0, m.start() - 80)
                ctx_end   = min(len(lower), m.end() + 80)
                ctx = lower[ctx_start:ctx_end]

                is_vendor = any(
                    re.search(neg, ctx, re.IGNORECASE)
                    for neg in _INDUSTRY_VENDOR_CONTEXT
                )
                if is_vendor:
                    continue

                label = m.group(0).strip()
                if label not in confirmed:
                    confirmed.append(label)
                if len(fragments) < 2:
                    s = max(0, m.start() - 30)
                    e = min(len(text), m.end() + 50)
                    frag = "…" + text[s:e].replace("\n", " ").strip() + "…"
                    if frag not in fragments:
                        fragments.append(frag)

        if len(confirmed) >= min_matches:
            industries.append(industry)
            industry_evidence[industry] = fragments

    return industries, industry_evidence


# ---------------------------------------------------------------------------
# Executive Framing Analysis
# ---------------------------------------------------------------------------

def _analyze_executive_framing(text: str) -> tuple[str, int, list[str], list[str]]:
    """
    Returns (level, score 0–100, operational_signals, executive_signals).
    Level: operational / technical / managerial / executive / strategic
    """
    op_sigs, _ = _find_signals(_OPERATIONAL_LANG_PATTERNS, text, max_evidence=0)
    ex_sigs, _ = _find_signals(_EXECUTIVE_LANG_PATTERNS, text, max_evidence=0)

    op_count = len(op_sigs)
    ex_count = len(ex_sigs)
    total = op_count + ex_count

    if total == 0:
        return "technical", 30, [], []

    ex_ratio = ex_count / total
    score = int(ex_ratio * 100)

    if ex_ratio >= 0.70:
        level = "strategic"
    elif ex_ratio >= 0.50:
        level = "executive"
    elif ex_ratio >= 0.35:
        level = "managerial"
    elif ex_ratio >= 0.18:
        level = "technical"
    else:
        level = "operational"

    return level, score, op_sigs, ex_sigs


# ---------------------------------------------------------------------------
# Positioning levels
# ---------------------------------------------------------------------------

def _rule_based_levels(profile: ResumeProfile) -> None:
    ai_count = len(profile.ai_signals)
    profile.ai_positioning_level = (
        PositioningLevel.HIGH   if ai_count >= 5 else
        PositioningLevel.MEDIUM if ai_count >= 2 else
        PositioningLevel.LOW
    )

    tf_count = len(profile.transformation_experience)
    profile.transformation_positioning_level = (
        PositioningLevel.HIGH   if tf_count >= 4 else
        PositioningLevel.MEDIUM if tf_count >= 2 else
        PositioningLevel.LOW
    )

    exec_ok = (
        profile.leadership_scope in ("company", "division")
        or profile.organizational_scale in ("enterprise", "holding")
    )
    profile.executive_positioning_level = (
        PositioningLevel.HIGH if exec_ok else PositioningLevel.MEDIUM
    )

    ops_count = len(profile.operational_scope)
    profile.operational_depth_level = (
        PositioningLevel.HIGH   if ops_count >= 3 else
        PositioningLevel.MEDIUM if ops_count >= 1 else
        PositioningLevel.LOW
    )

    # Strategic visibility — expanded signal set
    sv_labels, sv_ev = _find_signals(_STRATEGIC_VISIBILITY_PATTERNS, profile.raw_text)
    profile.business_signals = sv_labels
    profile.evidence["business_signals"] = sv_ev
    sv_count = len(sv_labels)
    profile.strategic_visibility_level = (
        PositioningLevel.HIGH   if sv_count >= 6 else
        PositioningLevel.MEDIUM if sv_count >= 3 else
        PositioningLevel.LOW
    )


# ---------------------------------------------------------------------------
# Main extraction
# ---------------------------------------------------------------------------

def _extract_rule_based(text: str, source_file: str, source_type: str) -> ResumeProfile:
    profile = ResumeProfile(raw_text=text, source_file=source_file, source_type=source_type)

    profile.years_experience = _estimate_years(text)
    profile.roles_history = parse_roles_history(text)

    # Composite scale
    scale, scale_signals, scale_rationale = _detect_scale_composite(text)
    profile.organizational_scale          = scale
    profile.organizational_scale_signals  = scale_signals
    profile.organizational_scale_rationale = scale_rationale
    profile.leadership_scope              = _detect_leadership_scope(text)

    # AI signals
    ai_labels, ai_ev = _find_signals(_AI_PATTERNS, text)
    profile.ai_signals = ai_labels
    profile.evidence["ai_signals"] = ai_ev

    # Transformation
    tf_labels, tf_ev = _find_signals(_TRANSFORM_PATTERNS, text)
    profile.transformation_experience = tf_labels
    profile.evidence["transformation"] = tf_ev

    # Governance
    gov_labels, gov_ev = _find_signals(_GOVERNANCE_PATTERNS, text)
    profile.governance_signals = gov_labels
    profile.evidence["governance"] = gov_ev

    # Technology domains
    domains: list[str] = []
    domain_ev: list[str] = []
    for domain, patterns in _TECH_PATTERNS.items():
        labels, ev = _find_signals(patterns, text, max_evidence=1)
        if labels:
            domains.append(domain)
            domain_ev.extend(ev[:1])
    profile.technology_domains = domains
    profile.evidence["technology_domains"] = domain_ev

    # Industries (v2 — contextual)
    profile.industries, industry_ev = _detect_industries_v2(text)
    profile.evidence["industries"] = [
        f for frags in industry_ev.values() for f in frags
    ][:6]

    # Operational scope
    ops_labels, ops_ev = _find_signals(_OPERATIONAL_SCOPE_PATTERNS, text)
    profile.operational_scope = ops_labels
    profile.evidence["operational_scope"] = ops_ev

    # Innovation
    inn_labels, _ = _find_signals(_INNOVATION_PATTERNS, text)
    profile.innovation_signals = inn_labels

    # Executive framing analysis
    framing_level, framing_score, op_lang, ex_lang = _analyze_executive_framing(text)
    profile.executive_framing_level      = framing_level
    profile.executive_framing_score      = framing_score
    profile.operational_language_signals = op_lang
    profile.executive_language_signals   = ex_lang

    _rule_based_levels(profile)
    return profile


# ---------------------------------------------------------------------------
# LLM enrichment
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
Ты — эксперт по executive career intelligence и resume analysis для топ-менеджеров уровня CIO/CTO/CDTO/CAIO.
Анализируешь резюме и строишь структурированный intelligence profile.
Отвечай строго в формате JSON без лишнего текста. Все строковые значения — на русском языке."""

_LLM_TEMPLATE = """\
Резюме:
---
{resume_text}
---

Контекст rule-based анализа:
- AI signals: {ai_count} · Трансформация: {tf_count} · Домены: {domains}
- Стаж: ~{years} лет · Scope: {scope} · Scale: {scale}
- Executive framing (rule-based): {framing_level} (score {framing_score}/100)

Верни ТОЛЬКО валидный JSON (без markdown):
{{
  "executive_summary": "2-3 предложения: кто этот человек профессионально",
  "strategic_positioning": "1 sentence executive positioning statement",
  "professional_archetype": "CIO|CDTO|CAIO|Transformation Executive|Operations Executive|AI Transformation Leader",
  "trajectory_markers": ["3-5 карьерных траекторий"],
  "strongest_assets": ["топ-5 executive strengths с опорой на конкретный опыт"],
  "weak_visibility_areas": [
    "Конкретная область + что скрыто + scale/ROI/governance/ownership который не показан"
  ],
  "executive_reframing_opportunities": [
    {{
      "original": "точная фраза из CV (операционная формулировка)",
      "suggested": "executive framing — тот же опыт в стратегическом языке",
      "reason": "почему это важно для C-level позиционирования"
    }}
  ],
  "ai_positioning_level": "LOW|MEDIUM|HIGH",
  "transformation_positioning_level": "LOW|MEDIUM|HIGH",
  "executive_positioning_level": "LOW|MEDIUM|HIGH",
  "operational_depth_level": "LOW|MEDIUM|HIGH",
  "strategic_visibility_level": "LOW|MEDIUM|HIGH"
}}

Требования:
- weak_visibility_areas: НЕ просто "AI weak". Формат: "AI initiatives описаны без ROI, scale, governance и business impact — hidden CAIO-potential"
- executive_reframing_opportunities: 3 конкретных примера из CV с оригинальной фразой и executive reframe
- trajectory_markers: конкретные trajectory paths (CIO / CDTO / CAIO / Enterprise Architecture / и т.д.)
- strongest_assets: конкретный опыт, а не абстракции"""


def _enrich_with_llm(profile: ResumeProfile, text: str) -> None:
    api_key = _get_api_key()
    if not api_key:
        _fill_fallback(profile)
        return

    try:
        prompt = _LLM_TEMPLATE.format(
            resume_text=text[:4000],
            ai_count=len(profile.ai_signals),
            tf_count=len(profile.transformation_experience),
            domains=", ".join(profile.technology_domains[:6]) or "—",
            years=profile.years_experience,
            scope=profile.leadership_scope,
            scale=profile.organizational_scale,
            framing_level=profile.executive_framing_level,
            framing_score=profile.executive_framing_score,
        )

        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=3000,
            system=_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
        )

        raw = response.content[0].text.strip()
        if raw.startswith("```"):
            lines = raw.split("\n")
            raw = "\n".join(lines[1:])
            if raw.endswith("```"):
                raw = raw[: raw.rfind("```")].rstrip()
        start, end = raw.find("{"), raw.rfind("}") + 1
        if start >= 0 and end > start:
            raw = raw[start:end]
        data = json.loads(raw)
        _apply_llm(profile, data)

    except Exception:
        _fill_fallback(profile)


def _apply_llm(profile: ResumeProfile, data: dict[str, Any]) -> None:
    profile.executive_summary      = data.get("executive_summary", "")
    profile.strategic_positioning  = data.get("strategic_positioning", "")
    profile.professional_archetype = data.get("professional_archetype", "")
    profile.trajectory_markers     = data.get("trajectory_markers") or []
    profile.strongest_assets       = data.get("strongest_assets") or []
    profile.weak_visibility_areas  = data.get("weak_visibility_areas") or []
    profile.executive_reframing_opportunities = data.get("executive_reframing_opportunities") or []

    def _level(key: str) -> PositioningLevel:
        raw = data.get(key, "")
        return PositioningLevel(raw) if raw in PositioningLevel._value2member_map_ else PositioningLevel.MEDIUM

    profile.ai_positioning_level             = _level("ai_positioning_level")
    profile.transformation_positioning_level = _level("transformation_positioning_level")
    profile.executive_positioning_level      = _level("executive_positioning_level")
    profile.operational_depth_level          = _level("operational_depth_level")
    profile.strategic_visibility_level       = _level("strategic_visibility_level")


def _fill_fallback(profile: ResumeProfile) -> None:
    domains = ", ".join(profile.technology_domains[:3]) or "ИТ-управление"
    years   = profile.years_experience

    profile.executive_summary = (
        f"Опытный ИТ-руководитель с ~{years}-летним стажем в областях: {domains}. "
        "Детальный semantic analysis требует подключения LLM."
    )
    profile.strategic_positioning  = f"Executive IT leader с опытом в {domains}."
    profile.professional_archetype = "CIO"
    if not profile.trajectory_markers:
        profile.trajectory_markers = ["CIO", "ИТ-директор", "Digital Leader"]
    if not profile.strongest_assets:
        profile.strongest_assets = [f"Опыт в {d}" for d in profile.technology_domains[:5]] or ["ИТ-управление"]
    if not profile.weak_visibility_areas:
        profile.weak_visibility_areas = ["Требуется LLM для анализа narrative gaps"]
    if not profile.executive_reframing_opportunities:
        profile.executive_reframing_opportunities = []
    profile.confidence_notes.append("LLM недоступен — использован rule-based fallback")


def _get_api_key() -> str:
    env_path = Path(__file__).resolve().parent.parent / ".env"
    if env_path.exists():
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                os.environ.setdefault(k.strip(), v.strip())
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    return key if key and not key.startswith("sk-ant-ВАШ") else ""
